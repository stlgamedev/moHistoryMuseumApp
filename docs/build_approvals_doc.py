"""Builds a single Word doc listing every question in the game, grouped
by section, formatted for stakeholder review/approval.

Output: docs/Question_Approvals.docx
"""

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONTENT = ROOT / "public" / "content"
INDEX_PATH = CONTENT / "sections.json"
OUT = ROOT / "docs" / "Question_Approvals.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)

TIERS = ["youth", "teen", "adult"]
TIER_LABELS = {"youth": "Youth", "teen": "Teen", "adult": "Adult"}


def set_cell_bg(cell, hex_color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def kv_row(table, label, value=""):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    for run in row[0].paragraphs[0].runs:
        run.bold = True
    set_cell_bg(row[0], "EAF1F8")
    row[0].width = Cm(4.5)
    row[1].width = Cm(12)
    return row


def heading(doc, text, *, size, color=ACCENT, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = True
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def fmt_choices(choices):
    out = []
    for c in choices:
        mark = " *" if c.get("correct") else ""
        out.append(f"{c['label']}{mark}")
    return "  ·  ".join(out)


def fmt_scan(scan):
    v = scan["verification"]
    target = (
        "Words: " + ", ".join(v["anyOf"])
        if v["kind"] == "ocr"
        else "Object: " + ", ".join(v["classes"])
    )
    return f"{scan['prompt']}  →  {target}"


def fmt_variant(variant):
    if variant["type"] == "multi-image":
        return f"[picture] {variant['prompt']}\n{fmt_choices(variant['choices'])}"
    if variant["type"] == "multi-text":
        return f"[multiple choice] {variant['prompt']}\n{fmt_choices(variant['choices'])}"
    extras = (
        f" (also accepted: {', '.join(variant['acceptableAnswers'])})"
        if variant.get("acceptableAnswers")
        else ""
    )
    return f"[type-in] {variant['prompt']}\nAnswer: {variant['answer']}{extras}"


def question_table(doc, question):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False

    kv_row(t, "Topic", question.get("topic", question["id"]))
    if question.get("hintText"):
        kv_row(t, "Hint", question["hintText"])
    if question.get("hintImage"):
        kv_row(t, "Hint image", question["hintImage"])

    by_tier = {v["ageTier"]: v for v in question["variants"]}
    for tier in TIERS:
        v = by_tier.get(tier)
        if not v:
            kv_row(t, TIER_LABELS[tier], "— not provided —")
            continue
        kv_row(t, TIER_LABELS[tier], fmt_variant(v))
        if v.get("scan"):
            kv_row(t, "  ↳ Camera (optional)", fmt_scan(v["scan"]))

    kv_row(t, "Approval", "☐ Approved    ☐ Changes (note below)    ☐ Rejected")
    kv_row(t, "Notes", "")


def section_block(doc, section):
    heading(doc, section["name"], size=16)
    if section.get("description"):
        p = doc.add_paragraph()
        r = p.add_run(section["description"])
        r.italic = True
        r.font.color.rgb = GREY
        r.font.size = Pt(10)
    if section.get("patchName"):
        p = doc.add_paragraph()
        r = p.add_run(f"Patch awarded: {section['patchName']}")
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    for i, question in enumerate(section["questions"]):
        if i > 0:
            doc.add_paragraph()
        question_table(doc, question)


def cover(doc, sections_data):
    heading(
        doc,
        "Question Approvals",
        size=22,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    heading(
        doc,
        "Missouri History Museum Scavenger Hunt",
        size=13,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
    caption(doc, f"Generated {date.today().isoformat()}")
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "Please review each question below. For each one, mark Approved, "
        "Changes (and add a note), or Rejected. Mark up directly in this "
        "document or print and hand back."
    )
    r.font.size = Pt(10)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("Sections in this game:")
    r.bold = True
    r.font.size = Pt(11)
    for s in sections_data:
        bullet = doc.add_paragraph(style="List Bullet")
        run = bullet.add_run(f"{s['name']}  ({len(s['questions'])} question(s))")
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(
        "Symbols used: * marks the correct multiple-choice answer.  "
        "[picture] = pick a picture · [multiple choice] = pick a word · "
        "[type-in] = type the answer.  Camera (optional) prompts ask the "
        "player to point their phone at a plaque or object; if the camera "
        "isn't available the player gets the regular question instead."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def main():
    index = json.loads(INDEX_PATH.read_text())
    sections_data = [json.loads((CONTENT / s["path"]).read_text()) for s in index["sections"]]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    cover(doc, sections_data)

    for section in sections_data:
        doc.add_page_break()
        section_block(doc, section)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
