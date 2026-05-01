"""Builds a 2-page Word doc:
  page 1: a worked example
  page 2: a blank, printable question form

The example pulls from the first question of the first section in
public/content/sections.json so it stays grounded in real content.

Output: resources/Question_Worksheet.docx
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONTENT = ROOT / "public" / "content"
INDEX_PATH = CONTENT / "sections.json"
OUT = ROOT / "docs" / "Question_Worksheet.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)

TIERS = ["youth", "teen", "adult"]
TIER_LABELS = {"youth": "Youth", "teen": "Teen", "adult": "Adult"}


def set_cell_bg(cell, hex_color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def two_col_table(doc):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    return t


def kv_row(table, label, value=""):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    for run in row[0].paragraphs[0].runs:
        run.bold = True
    set_cell_bg(row[0], "EAF1F8")
    row[0].width = Cm(4.5)
    row[1].width = Cm(12)
    return row


def title(doc, text, *, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = ACCENT


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def warning_banner(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cell, "C0392B")


def legend(doc):
    p = doc.add_paragraph()
    r = p.add_run(
        "Format A = pick a picture (3 choices, mark correct with *)   ·   "
        "Format B = pick a word (3 choices, mark correct with *)   ·   "
        "Format C = type the answer"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


# ---------- example page ----------

def fmt_choices(choices, with_images):
    out = []
    for c in choices:
        mark = " *" if c.get("correct") else ""
        if with_images and c.get("image"):
            out.append(f"{c['label']}{mark} (image: {c['image']})")
        else:
            out.append(f"{c['label']}{mark}")
    return "  ·  ".join(out)


def fmt_scan(scan):
    v = scan["verification"]
    target = (
        "Words: " + ", ".join(v["anyOf"])
        if v["kind"] == "ocr"
        else "Object: " + ", ".join(v["classes"])
    )
    return f"{scan['prompt']}  →  {target}"


def fmt_variant(variant):
    if variant["type"] == "multi-image":
        return f"[A] {variant['prompt']}\n{fmt_choices(variant['choices'], True)}"
    if variant["type"] == "multi-text":
        return f"[B] {variant['prompt']}\n{fmt_choices(variant['choices'], False)}"
    extras = (
        f" (also: {', '.join(variant['acceptableAnswers'])})"
        if variant.get("acceptableAnswers")
        else ""
    )
    return f"[C] {variant['prompt']}\nAnswer: {variant['answer']}{extras}"


def example_page(doc, section, question):
    warning_banner(doc, "EXAMPLE — DO NOT FILL OUT.  Use page 2 (Question Worksheet) for new questions.")
    doc.add_paragraph()
    title(doc, "Example", size=18)
    caption(doc, "Here's how a finished question looks. The next page is the blank to print and fill in.")
    doc.add_paragraph()

    t = two_col_table(doc)
    kv_row(t, "Section", section["name"])
    kv_row(t, "Topic", question.get("topic", question["id"]))
    if question.get("hintText"):
        kv_row(t, "Hint", question["hintText"])
    if question.get("hintImage"):
        kv_row(t, "Hint image", question["hintImage"])

    by_tier = {v["ageTier"]: v for v in question["variants"]}
    for tier in TIERS:
        v = by_tier.get(tier)
        if not v:
            continue
        kv_row(t, TIER_LABELS[tier], fmt_variant(v))
        if v.get("scan"):
            kv_row(t, "  ↳ Camera (optional)", fmt_scan(v["scan"]))

    doc.add_paragraph()
    legend(doc)


# ---------- blank page ----------

def blank_page(doc, section_names):
    title(doc, "Question Worksheet", size=18)
    caption(doc, "Print as many copies as you need — one per question.")

    p = doc.add_paragraph()
    r = p.add_run(f"Sections to choose from: {', '.join(section_names)}.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    t = two_col_table(doc)
    kv_row(t, "Section")
    kv_row(t, "Topic")
    kv_row(t, "Hint")
    kv_row(t, "Hint image")
    for tier in TIERS:
        kv_row(t, TIER_LABELS[tier])
        kv_row(t, "  Format (A/B/C)")
        kv_row(t, "  Choices / answer")
        kv_row(t, "  Camera (optional)")

    doc.add_paragraph()
    legend(doc)


# ---------- top-level ----------

def main():
    index = json.loads(INDEX_PATH.read_text())
    sections_data = [json.loads((CONTENT / s["path"]).read_text()) for s in index["sections"]]
    section_names = [s["name"] for s in sections_data]
    first_section = sections_data[0]
    first_question = first_section["questions"][0]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    example_page(doc, first_section, first_question)
    doc.add_page_break()
    blank_page(doc, section_names)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
